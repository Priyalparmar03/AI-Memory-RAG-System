from __future__ import annotations
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional
from docx import Document
from .base_loader import BaseLoader, LoaderError

logger = logging.getLogger(__name__)


# ==========================================================
# DOCX Loader
# ==========================================================

class DOCXLoader(BaseLoader):
    SUPPORTED_EXTENSIONS = [
        ".docx",
    ]

    # ======================================================
    # Constructor
    # ======================================================

    def __init__(
        self,
        file_path: str,
    ):

        super().__init__(file_path)

        self.validate_extension(

            self.SUPPORTED_EXTENSIONS

        )

        self.document: Optional[
            Document
        ] = None

        logger.info(

            f"DOCXLoader initialized "

            f"for {self.file_name}"

        )

    # ======================================================
    # Open Document
    # ======================================================

    def open_document(
        self,
    ) -> Document:
        """
        Open DOCX document.
        """

        if self.document is None:

            try:

                self.document = Document(

                    self.file_path

                )

            except Exception as error:

                raise LoaderError(

                    f"Unable to open DOCX: "

                    f"{error}"

                )

        return self.document

    # ======================================================
    # Close Document
    # ======================================================

    def close_document(
        self,
    ) -> None:
        """
        Release document.
        """

        self.document = None

    # ======================================================
    # Extract Text
    # ======================================================

    def extract_text(
        self,
    ) -> str:
        """
        Extract complete text.
        """

        document = self.open_document()

        paragraphs = [

            paragraph.text

            for paragraph

            in document.paragraphs

            if paragraph.text.strip()

        ]

        return "\n".join(

            paragraphs

        )

    # ======================================================
    # Paragraphs
    # ======================================================

    def paragraphs(
        self,
    ) -> List[str]:
        """
        Return paragraph list.
        """

        document = self.open_document()

        return [

            paragraph.text

            for paragraph

            in document.paragraphs

        ]

    # ======================================================
    # Metadata
    # ======================================================

    def document_metadata(
        self,
    ) -> Dict[str, Any]:
        """
        Basic metadata.
        """

        document = self.open_document()

        properties = (

            document.core_properties

        )

        return {

            "title":

                properties.title,

            "author":

                properties.author,

            "subject":

                properties.subject,

            "category":

                properties.category,

            "keywords":

                properties.keywords,

            "comments":

                properties.comments,

            "created":

                properties.created,

            "modified":

                properties.modified,

            "last_modified_by":

                properties.last_modified_by,

            "revision":

                properties.revision,

            "paragraphs":

                len(

                    document.paragraphs

                ),

            "sections":

                len(

                    document.sections

                ),

        }

    # ======================================================
    # Statistics
    # ======================================================

    def statistics(
        self,
    ) -> Dict[str, Any]:
        """
        Document statistics.
        """

        text = self.extract_text()

        return {

            "characters":

                len(text),

            "words":

                len(

                    text.split()

                ),

            "lines":

                len(

                    text.splitlines()

                ),

            "paragraphs":

                len(

                    self.paragraphs()

                ),

        }

    # ======================================================
    # Load
    # ======================================================

    def load(
        self,
    ) -> Dict[str, Any]:
        """
        Load DOCX document.
        """

        text = self.extract_text()

        metadata = self.document_metadata()

        metadata.update(

            self.statistics()

        )

        return self.build_result(

            text=text,

            pages=[],

            tables=[],

            images=[],

            extra_metadata=metadata,

        )

    # ======================================================
    # Preview
    # ======================================================

    def preview(
        self,
        characters: int = 1000,
    ) -> str:
        """
        Preview document.
        """

        return self.extract_text()[:characters]

    # ======================================================
    # Empty Check
    # ======================================================

    def is_empty(
        self,
    ) -> bool:
        """
        Check empty document.
        """

        return (

            len(

                self.extract_text().strip()

            )

            == 0

        )

    # ======================================================
    # String Representation
    # ======================================================

    def __repr__(
        self,
    ):

        return (

            "DOCXLoader("

            f"file='{self.file_name}', "

            f"paragraphs={len(self.paragraphs())}"

            ")"

        )

# ======================================================
# Extract Headings
# ======================================================

def extract_headings(
    self,
) -> List[Dict[str, Any]]:
    """
    Extract document headings.
    """

    document = self.open_document()

    headings = []

    for paragraph in document.paragraphs:

        style = paragraph.style.name

        if style.startswith("Heading"):

            try:

                level = int(

                    style.split()[-1]

                )

            except Exception:

                level = 1

            headings.append(

                {

                    "level": level,

                    "text": paragraph.text,

                    "style": style,

                }

            )

    return headings


# ======================================================
# Extract Tables
# ======================================================

def extract_tables(
    self,
) -> List[List[List[str]]]:
    """
    Extract all tables.
    """

    document = self.open_document()

    tables = []

    for table in document.tables:

        rows = []

        for row in table.rows:

            rows.append(

                [

                    cell.text.strip()

                    for cell

                    in row.cells

                ]

            )

        tables.append(

            rows

        )

    return tables


# ======================================================
# Extract Images
# ======================================================

def extract_images(
    self,
) -> List[Dict[str, Any]]:
    """
    Extract image metadata.
    """

    document = self.open_document()

    images = []

    relationships = document.part.rels

    image_number = 0

    for rel in relationships.values():

        if "image" in rel.target_ref:

            image_number += 1

            images.append(

                {

                    "id": image_number,

                    "name":

                        Path(

                            rel.target_ref

                        ).name,

                    "path":

                        rel.target_ref,

                }

            )

    return images


# ======================================================
# Extract Hyperlinks
# ======================================================

def extract_hyperlinks(
    self,
) -> List[Dict[str, str]]:
    """
    Extract hyperlinks.
    """

    document = self.open_document()

    links = []

    relationships = document.part.rels

    for rel in relationships.values():

        if (

            "hyperlink"

            in rel.reltype

        ):

            links.append(

                {

                    "url":

                        rel.target_ref,

                }

            )

    return links


# ======================================================
# Extract Headers
# ======================================================

def extract_headers(
    self,
) -> List[str]:
    """
    Extract document headers.
    """

    document = self.open_document()

    headers = []

    for section in document.sections:

        header = section.header

        for paragraph in header.paragraphs:

            if paragraph.text.strip():

                headers.append(

                    paragraph.text

                )

    return headers


# ======================================================
# Extract Footers
# ======================================================

def extract_footers(
    self,
) -> List[str]:
    """
    Extract document footers.
    """

    document = self.open_document()

    footers = []

    for section in document.sections:

        footer = section.footer

        for paragraph in footer.paragraphs:

            if paragraph.text.strip():

                footers.append(

                    paragraph.text

                )

    return footers


# ======================================================
# Search Text
# ======================================================

def search(
    self,
    keyword: str,
) -> List[Dict[str, Any]]:
    """
    Search document.
    """

    keyword = keyword.lower()

    results = []

    for index, paragraph in enumerate(

        self.paragraphs(),

        start=1,

    ):

        if (

            keyword

            in

            paragraph.lower()

        ):

            results.append(

                {

                    "paragraph":

                        index,

                    "text":

                        paragraph,

                }

            )

    return results


# ======================================================
# Heading Count
# ======================================================

def heading_count(
    self,
) -> int:
    """
    Number of headings.
    """

    return len(

        self.extract_headings()

    )


# ======================================================
# Table Count
# ======================================================

def table_count(
    self,
) -> int:
    """
    Number of tables.
    """

    return len(

        self.extract_tables()

    )


# ======================================================
# Image Count
# ======================================================

def image_count(
    self,
) -> int:
    """
    Number of embedded images.
    """

    return len(

        self.extract_images()

    )


# ======================================================
# Hyperlink Count
# ======================================================

def hyperlink_count(
    self,
) -> int:
    """
    Number of hyperlinks.
    """

    return len(

        self.extract_hyperlinks()

    )

# ======================================================
# Extract Sections
# ======================================================

def extract_sections(
    self,
) -> List[Dict[str, Any]]:
    """
    Extract document sections.
    """

    document = self.open_document()

    sections = []

    for index, section in enumerate(

        document.sections,

        start=1,

    ):

        sections.append(

            {

                "section": index,

                "start_type":

                    str(

                        section.start_type

                    ),

                "orientation":

                    str(

                        section.orientation

                    ),

                "page_width":

                    section.page_width,

                "page_height":

                    section.page_height,

                "left_margin":

                    section.left_margin,

                "right_margin":

                    section.right_margin,

                "top_margin":

                    section.top_margin,

                "bottom_margin":

                    section.bottom_margin,

            }

        )

    return sections


# ======================================================
# Extract Styles
# ======================================================

def extract_styles(
    self,
) -> List[str]:
    """
    Extract used paragraph styles.
    """

    document = self.open_document()

    styles = set()

    for paragraph in document.paragraphs:

        if paragraph.style:

            styles.add(

                paragraph.style.name

            )

    return sorted(

        list(styles)

    )


# ======================================================
# Extract Comments
# ======================================================

def extract_comments(
    self,
) -> List[Dict[str, Any]]:
    """
    Extract comments.

    NOTE:
    python-docx currently has
    limited support for comments.
    """

    logger.warning(

        "Comment extraction "

        "is not fully supported "

        "by python-docx."

    )

    return []


# ======================================================
# Core Properties
# ======================================================

def extract_core_properties(
    self,
) -> Dict[str, Any]:
    """
    Extract document core properties.
    """

    props = (

        self.open_document()

        .core_properties

    )

    return {

        "title":

            props.title,

        "author":

            props.author,

        "subject":

            props.subject,

        "category":

            props.category,

        "keywords":

            props.keywords,

        "language":

            props.language,

        "identifier":

            props.identifier,

        "revision":

            props.revision,

        "version":

            props.version,

        "created":

            props.created,

        "modified":

            props.modified,

        "last_modified_by":

            props.last_modified_by,

    }


# ======================================================
# Statistics
# ======================================================

def statistics(
    self,
) -> Dict[str, Any]:
    """
    Advanced document statistics.
    """

    text = self.extract_text()

    words = text.split()

    paragraphs = self.paragraphs()

    return {

        "characters":

            len(text),

        "characters_no_spaces":

            len(

                text.replace(

                    " ",

                    "",

                )

            ),

        "words":

            len(words),

        "lines":

            len(

                text.splitlines()

            ),

        "paragraphs":

            len(paragraphs),

        "headings":

            self.heading_count(),

        "tables":

            self.table_count(),

        "images":

            self.image_count(),

        "hyperlinks":

            self.hyperlink_count(),

        "sections":

            len(

                self.extract_sections()

            ),

        "styles":

            len(

                self.extract_styles()

            ),

    }


# ======================================================
# Preview Paragraph
# ======================================================

def preview_paragraph(
    self,
    index: int,
    characters: int = 500,
) -> str:
    """
    Preview one paragraph.
    """

    paragraphs = self.paragraphs()

    if (

        index < 0

        or

        index >= len(paragraphs)

    ):

        return ""

    return paragraphs[

        index

    ][:characters]


# ======================================================
# Summary
# ======================================================

def summary(
    self,
) -> Dict[str, Any]:
    """
    Human-readable summary.
    """

    metadata = self.document_metadata()

    stats = self.statistics()

    return {

        "file":

            self.file_name,

        "title":

            metadata.get(

                "title"

            ),

        "author":

            metadata.get(

                "author"

            ),

        "paragraphs":

            stats["paragraphs"],

        "headings":

            stats["headings"],

        "tables":

            stats["tables"],

        "images":

            stats["images"],

        "hyperlinks":

            stats["hyperlinks"],

        "sections":

            stats["sections"],

        "words":

            stats["words"],

    }


# ======================================================
# Export Structure
# ======================================================

def export_structure(
    self,
) -> Dict[str, Any]:
    """
    Export document structure.
    """

    return {

        "metadata":

            self.document_metadata(),

        "core_properties":

            self.extract_core_properties(),

        "paragraphs":

            self.paragraphs(),

        "headings":

            self.extract_headings(),

        "tables":

            self.extract_tables(),

        "images":

            self.extract_images(),

        "hyperlinks":

            self.extract_hyperlinks(),

        "headers":

            self.extract_headers(),

        "footers":

            self.extract_footers(),

        "sections":

            self.extract_sections(),

        "styles":

            self.extract_styles(),

        "statistics":

            self.statistics(),

    }

# ======================================================
# Diagnostics
# ======================================================

def diagnostics(
    self,
) -> Dict[str, Any]:
    """
    Document diagnostics.
    """

    return {

        "loader":

            self.__class__.__name__,

        "file":

            self.file_name,

        "extension":

            self.extension,

        "mime_type":

            self.mime_type,

        "file_size":

            self.file_size,

        "is_empty":

            self.is_empty(),

        "statistics":

            self.statistics(),

    }


# ======================================================
# Benchmark
# ======================================================

def benchmark(
    self,
) -> Dict[str, Any]:
    """
    Benchmark document loading.
    """

    import time

    start = time.perf_counter()

    text = self.extract_text()

    elapsed = (

        time.perf_counter()

        -

        start

    )

    words = len(

        text.split()

    )

    paragraphs = len(

        self.paragraphs()

    )

    return {

        "execution_time":

            round(

                elapsed,

                4,

            ),

        "words_per_second":

            round(

                words

                /

                max(

                    elapsed,

                    1e-9,

                ),

                2,

            ),

        "paragraphs_per_second":

            round(

                paragraphs

                /

                max(

                    elapsed,

                    1e-9,

                ),

                2,

            ),

    }


# ======================================================
# Validate Document
# ======================================================

def validate(
    self,
) -> bool:
    """
    Validate DOCX document.
    """

    try:

        self.open_document()

        return True

    except Exception:

        return False


# ======================================================
# Reload
# ======================================================

def reload(
    self,
) -> None:
    """
    Reload document.
    """

    self.close_document()

    self.open_document()


# ======================================================
# Export Plain Text
# ======================================================

def export_text(
    self,
    output_path: str,
) -> str:
    """
    Export document text.
    """

    with open(

        output_path,

        "w",

        encoding="utf-8",

    ) as file:

        file.write(

            self.extract_text()

        )

    return output_path


# ======================================================
# Export Metadata
# ======================================================

def export_metadata(
    self,
) -> Dict[str, Any]:
    """
    Export complete metadata.
    """

    metadata = self.metadata()

    metadata.update(

        self.document_metadata()

    )

    metadata.update(

        self.statistics()

    )

    return metadata


# ======================================================
# Cleanup
# ======================================================

def cleanup(
    self,
) -> None:
    """
    Cleanup resources.
    """

    self.close_document()

    logger.info(

        "DOCX resources released."

    )


# ======================================================
# Context Manager
# ======================================================

def __enter__(
    self,
):

    self.open_document()

    return self


def __exit__(
    self,
    exc_type,
    exc_value,
    traceback,
):

    self.cleanup()


# ======================================================
# Python Protocols
# ======================================================

def __len__(
    self,
):
    """
    Number of paragraphs.
    """

    return len(

        self.paragraphs()

    )


def __iter__(
    self,
):
    """
    Iterate over paragraphs.
    """

    return iter(

        self.paragraphs()

    )


def __getitem__(
    self,
    index: int,
):
    """
    Get paragraph by index.
    """

    paragraphs = self.paragraphs()

    return paragraphs[index]


def __contains__(
    self,
    keyword: str,
):
    """
    Check keyword existence.
    """

    return (

        keyword.lower()

        in

        self.extract_text().lower()

    )


def __repr__(
    self,
):
    """
    String representation.
    """

    stats = self.statistics()

    return (

        "DOCXLoader("

        f"file='{self.file_name}', "

        f"paragraphs={stats['paragraphs']}, "

        f"headings={stats['headings']}, "

        f"tables={stats['tables']}"

        ")"

    )
