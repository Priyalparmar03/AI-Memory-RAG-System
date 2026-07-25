"""
services/export_service.py

Production Export Service

Responsibilities
----------------
- Conversation Export
- Analytics Export
- Memory Export
- RAG Export
- Backup
- Reports
"""

from __future__ import annotations
import zipfile
import shutil
import os
import json
import logging
import re
import csv
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt

from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate

logger = logging.getLogger(__name__)


class ExportService:
    """
    Export Service

    Handles exporting conversations and reports into multiple formats.
    """

    def __init__(
        self,
        chat_service,
        memory_service,
        analytics_service,
        rag_service,
        export_directory: str = "./exports",
    ) -> None:
        self.chat = chat_service
        self.memory = memory_service
        self.analytics = analytics_service
        self.rag = rag_service

        self.export_directory = Path(export_directory)
        self.export_directory.mkdir(parents=True, exist_ok=True)

        logger.info("ExportService initialized at %s", self.export_directory.resolve())

    # ==========================================================
    # Internal Utilities
    # ==========================================================

    def _sanitize_filename(self, filename: str) -> str:
        """Remove invalid filename characters."""
        return re.sub(r"[^A-Za-z0-9_.-]", "_", filename)

    def create_folder(self, folder_name: str) -> Path:
        """Create export sub-folder if it doesn't exist."""
        folder = self.export_directory / folder_name
        folder.mkdir(parents=True, exist_ok=True)
        return folder

    # ==========================================================
    # JSON Export
    # ==========================================================

    def export_conversation_json(self, conversation_id: str) -> Path:
        """Export conversation as JSON."""
        conversation = self.chat.export_conversation(conversation_id)

        folder = self.create_folder("json")
        filename = folder / f"{self._sanitize_filename(conversation_id)}.json"

        with filename.open("w", encoding="utf-8") as file:
            json.dump(conversation, file, indent=4, ensure_ascii=False)

        logger.info("Conversation exported to JSON: %s", filename)
        return filename

    # ==========================================================
    # Markdown Export
    # ==========================================================

    def export_conversation_markdown(self, conversation_id: str) -> Path:
        """Export conversation as Markdown."""
        conversation = self.chat.export_conversation(conversation_id)

        folder = self.create_folder("markdown")
        filename = folder / f"{self._sanitize_filename(conversation_id)}.md"

        with filename.open("w", encoding="utf-8") as file:
            file.write("# Conversation Export\n\n")
            file.write(f"**Conversation ID:** {conversation_id}\n\n")
            file.write(f"**Exported:** {datetime.now().isoformat()}\n\n")
            file.write("---\n\n")

            for message in conversation.get("messages", []):
                role = message.get("role", "unknown")
                content = message.get("content", "")
                file.write(f"## {role.title()}\n\n")
                file.write(f"{content}\n\n")

        logger.info("Conversation exported to Markdown: %s", filename)
        return filename

    # ==========================================================
    # HTML Export
    # ==========================================================

    def export_conversation_html(self, conversation_id: str) -> Path:
        """Export conversation as HTML."""
        conversation = self.chat.export_conversation(conversation_id)

        folder = self.create_folder("html")
        filename = folder / f"{self._sanitize_filename(conversation_id)}.html"

        html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>Conversation Export</title>
<style>
body{{font-family:Arial,Helvetica,sans-serif;background:#f5f5f5;margin:40px;}}
.container{{max-width:900px;margin:auto;}}
.user{{background:#d9fdd3;padding:15px;margin:15px 0;border-radius:8px;}}
.assistant{{background:#ffffff;padding:15px;margin:15px 0;border-radius:8px;border:1px solid #dddddd;}}
.system{{background:#eeeeee;padding:15px;margin:15px 0;border-radius:8px;}}
.footer{{margin-top:40px;font-size:12px;color:gray;}}
</style>
</head>
<body>
<div class="container">
<h1>Conversation Export</h1>
<p><b>Conversation ID:</b> {escape(conversation_id)}</p>
<p><b>Exported:</b> {datetime.now().isoformat()}</p>
<hr>
"""

        for message in conversation.get("messages", []):
            role = message.get("role", "system")
            content = escape(message.get("content", ""))
            html += f"""
<div class="{role}">
<h3>{role.title()}</h3>
<p>{content}</p>
</div>
"""

        html += """
<div class="footer">Generated by AI Memory RAG System</div>
</div>
</body>
</html>
"""

        with filename.open("w", encoding="utf-8") as file:
            file.write(html)

        logger.info("Conversation exported to HTML: %s", filename)
        return filename

    # ==========================================================
    # PDF Export
    # ==========================================================

    def export_conversation_pdf(self, conversation_id: str) -> Path:
        """Export conversation as PDF."""
        conversation = self.chat.export_conversation(conversation_id)

        folder = self.create_folder("pdf")
        filename = folder / f"{self._sanitize_filename(conversation_id)}.pdf"

        styles = getSampleStyleSheet()
        document = SimpleDocTemplate(str(filename))
        elements = []

        elements.append(Paragraph("<b>Conversation Export</b>", styles["Title"]))
        elements.append(Paragraph(f"Conversation ID: {conversation_id}", styles["Normal"]))
        elements.append(Paragraph(f"Exported: {datetime.now()}", styles["Normal"]))

        for message in conversation.get("messages", []):
            role = escape(message.get("role", "Unknown"))
            content = escape(message.get("content", ""))
            elements.append(Paragraph(f"<b>{role.title()}</b>", styles["Heading2"]))
            elements.append(Paragraph(content, styles["BodyText"]))

        document.build(elements)

        logger.info("PDF exported: %s", filename)
        return filename

    # ==========================================================
    # DOCX Export
    # ==========================================================

    def export_conversation_docx(self, conversation_id: str) -> Path:
        """Export conversation as DOCX."""
        conversation = self.chat.export_conversation(conversation_id)

        folder = self.create_folder("docx")
        filename = folder / f"{self._sanitize_filename(conversation_id)}.docx"

        document = Document()
        document.add_heading("Conversation Export", level=1)
        document.add_paragraph(f"Conversation ID: {conversation_id}")
        document.add_paragraph(f"Exported: {datetime.now()}")
        document.add_paragraph()

        for message in conversation.get("messages", []):
            role = message.get("role", "Unknown")
            content = message.get("content", "")

            heading = document.add_heading(role.title(), level=2)
            heading.style.font.size = Pt(14)

            document.add_paragraph(content)

        document.save(filename)

        logger.info("DOCX exported: %s", filename)
        return filename

    # ==========================================================
    # CSV Export
    # ==========================================================

    def export_conversation_csv(self, conversation_id: str) -> Path:
        """Export conversation as CSV."""
        conversation = self.chat.export_conversation(conversation_id)

        folder = self.create_folder("csv")
        filename = folder / f"{self._sanitize_filename(conversation_id)}.csv"

        with filename.open("w", newline="", encoding="utf-8") as csvfile:
            writer = csv.writer(csvfile)
            writer.writerow(["role", "content"])

            for message in conversation.get("messages", []):
                writer.writerow([message.get("role"), message.get("content")])

        logger.info("CSV exported: %s", filename)
        return filename

    # ==========================================================
    # Export Multiple Formats
    # ==========================================================

    def export_multiple_formats(self, conversation_id: str) -> Dict[str, str]:
        """Export a conversation into all supported formats."""
        exports = {
            "json": str(self.export_conversation_json(conversation_id)),
            "markdown": str(self.export_conversation_markdown(conversation_id)),
            "html": str(self.export_conversation_html(conversation_id)),
            "pdf": str(self.export_conversation_pdf(conversation_id)),
            "docx": str(self.export_conversation_docx(conversation_id)),
            "csv": str(self.export_conversation_csv(conversation_id)),
        }

        logger.info("Conversation exported in all formats.")
        return exports

    # ==========================================================
    # Analytics Export
    # ==========================================================

    def export_analytics_json(self) -> Path:
        """Export analytics dashboard."""
        dashboard = self.analytics.dashboard()

        folder = self.create_folder("analytics")
        filename = folder / "analytics_dashboard.json"

        with filename.open("w", encoding="utf-8") as file:
            json.dump(dashboard, file, indent=4, ensure_ascii=False)

        logger.info("Analytics exported.")
        return filename

    # ==========================================================
    # Memory Export
    # ==========================================================

    def export_memory_json(self, user_id: str) -> Path:
        """Export all long-term memories for a user."""
        memories = self.memory.export_memories(user_id)

        folder = self.create_folder("memory")
        filename = folder / f"{self._sanitize_filename(user_id)}_memory.json"

        with filename.open("w", encoding="utf-8") as file:
            json.dump(memories, file, indent=4, ensure_ascii=False)

        logger.info("Memory exported: %s", filename)
        return filename

    # ==========================================================
    # User Profile Export
    # ==========================================================

    def export_profile_json(self, user_id: str) -> Path:
        """Export user profile."""
        profile = self.memory.get_user_profile(user_id)

        folder = self.create_folder("profile")
        filename = folder / f"{self._sanitize_filename(user_id)}_profile.json"

        with filename.open("w", encoding="utf-8") as file:
            json.dump(profile, file, indent=4, ensure_ascii=False)

        logger.info("Profile exported.")
        return filename

    # ==========================================================
    # RAG Search Export
    # ==========================================================

    def export_rag_results(self, query: str) -> Path:
        """Export semantic search results."""
        results = self.rag.search(query=query)

        folder = self.create_folder("rag")
        filename = folder / "rag_results.json"

        with filename.open("w", encoding="utf-8") as file:
            json.dump(results, file, indent=4, ensure_ascii=False)

        logger.info("RAG results exported.")
        return filename

    # ==========================================================
    # Citation Export
    # ==========================================================

    def export_citations(self, query: str) -> Path:
        """Export retrieved citations."""
        results = self.rag.search(query)

        citations = [
            {
                "document": item.get("document"),
                "score": item.get("score"),
                "metadata": item.get("metadata"),
            }
            for item in results
        ]

        folder = self.create_folder("citations")
        filename = folder / "citations.json"

        with filename.open("w", encoding="utf-8") as file:
            json.dump(citations, file, indent=4, ensure_ascii=False)

        logger.info("Citation export complete.")
        return filename

    # ==========================================================
    # User Backup
    # ==========================================================

    def export_user_backup(self, user_id: str) -> Dict[str, Path]:
        """Export complete user data."""
        backup = {
            "memory": self.export_memory_json(user_id),
            "profile": self.export_profile_json(user_id),
        }

        logger.info("User backup created.")
        return backup

    # ==========================================================
    # ZIP Archive
    # ==========================================================

    def create_backup_zip(self, user_id: str) -> Path:
        """Create ZIP archive of exported data."""
        exports = self.export_user_backup(user_id)

        folder = self.create_folder("backup")
        zip_path = folder / f"{self._sanitize_filename(user_id)}_backup.zip"

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
            for file in exports.values():
                archive.write(file, arcname=file.name)

        logger.info("Backup ZIP created.")
        return zip_path

    # ==========================================================
    # Complete Backup
    # ==========================================================

    def export_everything(self, conversation_id: str, user_id: str) -> Dict[str, str]:
        """Export all supported resources."""
        data = {}
        data.update(self.export_multiple_formats(conversation_id))
        data["memory"] = str(self.export_memory_json(user_id))
        data["profile"] = str(self.export_profile_json(user_id))
        data["backup_zip"] = str(self.create_backup_zip(user_id))
        return data

    # ==========================================================
    # Restore Backup
    # ==========================================================

    def restore_backup(self, zip_file: str, destination: Optional[str] = None) -> Path:
        """Restore exported backup ZIP."""
        zip_path = Path(zip_file)

        if not zip_path.exists():
            raise FileNotFoundError(f"Backup not found: {zip_file}")

        if destination is None:
            destination = self.export_directory / "restored"

        destination = Path(destination)
        destination.mkdir(parents=True, exist_ok=True)

        with zipfile.ZipFile(zip_path, "r") as archive:
            archive.extractall(destination)

        logger.info("Backup restored to %s", destination)
        return destination

    # ==========================================================
    # Cleanup Exports
    # ==========================================================

    def cleanup_exports(self, older_than_days: int = 30) -> int:
        """Delete exports older than the specified number of days."""
        deleted = 0
        cutoff = datetime.now().timestamp() - (older_than_days * 86400)

        for file in self.export_directory.rglob("*"):
            if not file.is_file():
                continue

            if file.stat().st_mtime < cutoff:
                try:
                    file.unlink()
                    deleted += 1
                except Exception:
                    logger.exception("Unable to delete %s", file)

        logger.info("Cleanup complete (%d files removed)", deleted)
        return deleted

    # ==========================================================
    # Storage Statistics
    # ==========================================================

    def storage_statistics(self) -> Dict[str, Any]:
        """Storage usage statistics."""
        total_files = 0
        total_size = 0
        by_extension: Dict[str, Dict[str, int]] = {}

        for file in self.export_directory.rglob("*"):
            if not file.is_file():
                continue

            total_files += 1
            size = file.stat().st_size
            total_size += size

            extension = file.suffix.lower() or "no_extension"
            by_extension.setdefault(extension, {"count": 0, "size": 0})
            by_extension[extension]["count"] += 1
            by_extension[extension]["size"] += size

        return {
            "total_files": total_files,
            "total_size_bytes": total_size,
            "total_size_mb": round(total_size / (1024 * 1024), 2),
            "by_extension": by_extension,
        }

    # ==========================================================
    # Export Management
    # ==========================================================

    def list_exports(self) -> List[Dict[str, Any]]:
        """List all exported files."""
        exports: List[Dict[str, Any]] = []

        for file in self.export_directory.rglob("*"):
            if file.is_file():
                exports.append(
                    {
                        "name": file.name,
                        "path": str(file),
                        "size": file.stat().st_size,
                        "created": datetime.fromtimestamp(file.stat().st_ctime).isoformat(),
                    }
                )

        return sorted(exports, key=lambda x: x["created"], reverse=True)

    def delete_export(self, filename: str) -> bool:
        """Delete exported file."""
        path = Path(filename)

        if path.exists() and path.is_file():
            path.unlink()
            logger.info("Deleted export: %s", filename)
            return True

        logger.warning("Export not found: %s", filename)
        return False

    # ==========================================================
    # Dashboard
    # ==========================================================

    def dashboard(self) -> Dict[str, Any]:
        """Export dashboard."""
        exports = self.list_exports()
        total_size = sum(item["size"] for item in exports)

        return {
            "total_exports": len(exports),
            "directory": str(self.export_directory),
            "total_size_bytes": total_size,
        }

    # ==========================================================
    # Diagnostics
    # ==========================================================

    def diagnostics(self) -> Dict[str, Any]:
        """Service diagnostics."""
        return {
            "service": "ExportService",
            "status": "healthy",
            "export_directory": str(self.export_directory.resolve()),
            "directory_exists": self.export_directory.exists(),
            "directory_writable": os.access(self.export_directory, os.W_OK),
            "dashboard": self.dashboard(),
            "storage": self.storage_statistics(),
        }

    # ==========================================================
    # Reset Service
    # ==========================================================

    def reset(self) -> None:
        """Remove every exported file."""
        if self.export_directory.exists():
            shutil.rmtree(self.export_directory)

        self.export_directory.mkdir(parents=True, exist_ok=True)
        logger.warning("Export directory reset.")

    # ==========================================================
    # Service Information
    # ==========================================================

    def info(self) -> Dict[str, Any]:
        """Service metadata."""
        return {
            "service": "ExportService",
            "version": "1.0.0",
            "formats": ["json", "markdown", "html", "pdf", "docx", "csv", "zip"],
            "directory": str(self.export_directory),
        }

    # ==========================================================
    # Health
    # ==========================================================

    def health(self) -> Dict[str, Any]:
        """Service health."""
        healthy = self.export_directory.exists() and os.access(self.export_directory, os.W_OK)

        return {
            "status": "healthy" if healthy else "unhealthy",
            "directory": str(self.export_directory),
            "dashboard": self.dashboard(),
            "storage": self.storage_statistics(),
        }

    # ==========================================================
    # Shutdown
    # ==========================================================

    def close(self) -> None:
        """Gracefully shutdown ExportService."""
        logger.info("Closing ExportService...")
        logger.info("ExportService shutdown complete.")
